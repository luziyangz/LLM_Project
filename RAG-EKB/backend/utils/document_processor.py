import re
import os
from typing import List, Dict, Any
from datetime import datetime

# 在文件顶部添加导入
from .pdf_processor import EnhancedPDFProcessor
from .docx_processor import EnhancedDocxProcessor

class DocumentProcessor:
    """文档处理工具类，支持多种文件类型的处理"""
    
    def __init__(self, upload_dir: str = "backend/file"):
        self.upload_dir = upload_dir
        self.pdf_processor = EnhancedPDFProcessor()  # 添加PDF处理器
        self.docx_processor = EnhancedDocxProcessor()  # 添加DOCX处理器
        self.supported_types = {
            '.txt': self._process_txt,
            '.pdf': self._process_pdf_enhanced,  # 更新PDF处理方法
            '.docx': self._process_docx_enhanced,  # 更新DOCX处理方法
            '.md': self._process_markdown
        }
    
    def _process_pdf_enhanced(self, filepath: str) -> str:
        """增强的PDF处理方法"""
        try:
            # 使用增强处理器
            result = self.pdf_processor.process_pdf_comprehensive(filepath)
            
            # 将所有元素组合成文本
            content_parts = []
            
            for element in result['elements']:
                if element.element_type == 'header':
                    content_parts.append(f"\n## {element.content}\n")
                elif element.element_type == 'text':
                    content_parts.append(element.content)
                elif element.element_type == 'table':
                    content_parts.append(f"\n[表格]\n{element.content}\n")
                elif element.element_type == 'image':
                    content_parts.append(f"\n{element.content}\n")
            
            return self._clean_text("\n".join(content_parts))
            
        except Exception as e:
            # 降级到简单处理
            return self._process_pdf_simple(filepath)
    
    def _process_pdf_simple(self, filepath: str) -> str:
        """简单PDF处理（降级方案）"""
        try:
            import PyPDF2
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                for page in pdf_reader.pages:
                    content += page.extract_text()
                return self._clean_text(content)
        except Exception as e:
            raise Exception(f"处理PDF文件时出错: {str(e)}")
    
    def is_supported(self, filename: str) -> bool:
        """检查文件类型是否支持"""
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.supported_types
    
    def process_document(self, filepath: str) -> Dict[str, Any]:
        """处理文档并返回处理结果"""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"文件不存在: {filepath}")
        
        filename = os.path.basename(filepath)
        ext = os.path.splitext(filename)[1].lower()
        
        if ext not in self.supported_types:
            raise ValueError(f"不支持的文件类型: {ext}")
        
        # 调用对应的处理方法
        processor = self.supported_types[ext]
        content = processor(filepath)
        
        # 返回处理结果
        return {
            'filename': filename,
            'filepath': filepath,
            'file_type': ext,
            'content': content,
            'word_count': len(content),
            'processed_at': datetime.now().isoformat()
        }
    
    def _process_txt(self, filepath: str) -> str:
        """处理TXT文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(filepath, 'r', encoding=encoding) as f:
                        content = f.read()
                    # 清理文本
                    content = self._clean_text(content)
                    return content
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("无法解码文件，请检查文件编码")
            
        except Exception as e:
            raise Exception(f"处理TXT文件时出错: {str(e)}")
    
    def _process_pdf(self, filepath: str) -> str:
        """处理PDF文件（预留接口）"""
        # TODO: 实现PDF处理逻辑
        # 可以使用 PyPDF2, pdfplumber 等库
        raise NotImplementedError("PDF处理功能尚未实现")
    
    def _process_docx_enhanced(self, filepath: str) -> str:
        """增强的DOCX处理方法"""
        try:
            # 使用增强处理器
            result = self.docx_processor.process_docx_comprehensive(filepath)
            
            # 将所有元素组合成文本
            content_parts = []
            
            for element in result['elements']:
                if element.element_type == 'header':
                    # 根据样式名称确定标题级别
                    level = self._extract_heading_level(element.metadata.get('style_name', ''))
                    content_parts.append(f"\n{'#' * level} {element.content}\n")
                elif element.element_type == 'title':
                    content_parts.append(f"\n# {element.content}\n")
                elif element.element_type == 'text':
                    content_parts.append(element.content)
                elif element.element_type == 'table':
                    content_parts.append(f"\n{element.content}\n")
                elif element.element_type == 'list':
                    content_parts.append(f"- {element.content}")
                elif element.element_type in ['header', 'footer']:
                    content_parts.append(f"\n[{element.metadata['source']}] {element.content}\n")
            
            return self._clean_text("\n".join(content_parts))
            
        except Exception as e:
            # 降级到简单处理
            return self._process_docx_simple(filepath)
    
    def _extract_heading_level(self, style_name: str) -> int:
        """从样式名称提取标题级别"""
        if 'Heading' in style_name:
            try:
                level = int(style_name.replace('Heading ', '').strip())
                return min(level, 6)  # 最多6级标题
            except:
                pass
        return 1
    
    def _process_docx_simple(self, filepath: str) -> str:
        """简单DOCX处理（降级方案）"""
        try:
            from docx import Document
            
            doc = Document(filepath)
            content_parts = []
            
            # 简单提取所有段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            return self._clean_text("\n".join(content_parts))
            
        except ImportError:
            raise Exception("请安装 python-docx 库: pip install python-docx")
        except Exception as e:
            raise Exception(f"处理DOCX文件时出错: {str(e)}")
    
    def _extract_table_content(self, table) -> str:
        """提取表格内容"""
        try:
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_data.append(cell_text)
                
                if any(cell.strip() for cell in row_data):  # 跳过空行
                    table_data.append(' | '.join(row_data))
            
            return '\n'.join(table_data)
            
        except Exception as e:
            return f"[表格解析错误: {str(e)}]"
    
    def _extract_header_footer(self, doc) -> str:
        """提取页眉页脚内容"""
        try:
            header_footer_parts = []
            
            # 提取页眉
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            header_footer_parts.append(f"[页眉] {text}")
                
                # 提取页脚
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            header_footer_parts.append(f"[页脚] {text}")
            
            return '\n'.join(header_footer_parts) if header_footer_parts else ""
            
        except Exception:
            return ""
    
    def _process_markdown(self, filepath: str) -> str:
        """处理Markdown文件（预留接口）"""
        # TODO: 实现Markdown处理逻辑
        raise NotImplementedError("Markdown处理功能尚未实现")
    
    def _clean_text(self, text: str) -> str:
        """清理文本内容"""
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        # 移除首尾空白
        text = text.strip()
        return text
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """将文本分块（为后续向量化准备）"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # 如果不是最后一块，尝试在句号处分割
            if end < len(text):
                # 寻找最近的句号
                last_period = text.rfind('。', start, end)
                if last_period > start:
                    end = last_period + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # 设置下一块的起始位置（考虑重叠）
            start = end - overlap if end < len(text) else end
        
        return chunks
    
    def search_in_document(self, filepath: str, query: str) -> List[Dict[str, Any]]:
        """在文档中搜索关键词"""
        try:
            result = self.process_document(filepath)
            content = result['content']
            
            # 简单的关键词搜索
            matches = []
            lines = content.split('\n')
            
            for i, line in enumerate(lines):
                if query.lower() in line.lower():
                    matches.append({
                        'line_number': i + 1,
                        'content': line.strip(),
                        'filename': result['filename']
                    })
            
            return matches
            
        except Exception as e:
            raise Exception(f"搜索文档时出错: {str(e)}")