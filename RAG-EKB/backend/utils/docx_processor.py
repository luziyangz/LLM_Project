from typing import Dict, List, Any, Optional
from dataclasses import dataclass
import logging
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

@dataclass
class DocxElement:
    """DOCX元素数据类"""
    element_type: str  # 'text', 'header', 'table', 'footer', 'list'
    content: str
    section_number: int
    position: int  # 在文档中的位置
    metadata: Dict[str, Any]

class EnhancedDocxProcessor:
    """增强的DOCX处理器，支持文本、表格、标题、列表等结构化提取"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
        self.logger = logging.getLogger(__name__)
    
    def process_docx_comprehensive(self, filepath: str) -> Dict[str, Any]:
        """综合处理DOCX文档"""
        elements = []
        
        try:
            doc = Document(filepath)
            
            # 提取文档统计信息
            doc_stats = self._get_document_statistics(doc)
            
            # 处理文档中的所有元素
            position = 0
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # 处理段落
                    paragraph_elements = self._extract_paragraph_elements(element, doc, position)
                    elements.extend(paragraph_elements)
                    position += len(paragraph_elements)
                    
                elif isinstance(element, CT_Tbl):
                    # 处理表格
                    table_element = self._extract_table_element(element, doc, position)
                    if table_element:
                        elements.append(table_element)
                        position += 1
            
            # 处理页眉页脚
            header_footer_elements = self._extract_header_footer_elements(doc, position)
            elements.extend(header_footer_elements)
            
            return {
                'elements': elements,
                'document_stats': doc_stats,
                'processing_method': 'comprehensive_docx',
                'total_elements': len(elements)
            }
            
        except Exception as e:
            self.logger.error(f"处理DOCX文件失败: {str(e)}")
            raise Exception(f"处理DOCX文件时出错: {str(e)}")
    
    def _get_document_statistics(self, doc: Document) -> Dict[str, Any]:
        """获取文档统计信息"""
        stats = {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'total_sections': len(doc.sections),
            'has_headers': False,
            'has_footers': False,
            'has_images': False
        }
        
        # 检查页眉页脚
        for section in doc.sections:
            if section.header.paragraphs:
                stats['has_headers'] = True
            if section.footer.paragraphs:
                stats['has_footers'] = True
        
        # 检查图片
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                stats['has_images'] = True
                break
        
        return stats
    
    def _extract_paragraph_elements(self, element, doc: Document, position: int) -> List[DocxElement]:
        """提取段落元素"""
        elements = []
        paragraph = Paragraph(element, doc)
        text = paragraph.text.strip()
        
        if text:
            # 判断段落类型
            element_type = self._classify_paragraph_type(paragraph)
            
            # 提取段落元数据
            metadata = {
                'style_name': paragraph.style.name,
                'font_size': self._get_font_size(paragraph),
                'is_bold': self._is_bold(paragraph),
                'is_italic': self._is_italic(paragraph),
                'alignment': str(paragraph.alignment) if paragraph.alignment else 'left',
                'char_count': len(text),
                'word_count': len(text.split())
            }
            
            elements.append(DocxElement(
                element_type=element_type,
                content=text,
                section_number=0,  # 可以根据需要计算实际section
                position=position,
                metadata=metadata
            ))
        
        return elements
    
    def _extract_table_element(self, element, doc: Document, position: int) -> Optional[DocxElement]:
        """提取表格元素"""
        table = Table(element, doc)
        table_content = self._table_to_structured_text(table)
        
        if table_content:
            metadata = {
                'rows': len(table.rows),
                'columns': len(table.columns) if table.rows else 0,
                'has_header_row': self._has_header_row(table),
                'table_style': table.style.name if table.style else 'Normal'
            }
            
            return DocxElement(
                element_type='table',
                content=table_content,
                section_number=0,
                position=position,
                metadata=metadata
            )
        
        return None
    
    def _extract_header_footer_elements(self, doc: Document, start_position: int) -> List[DocxElement]:
        """提取页眉页脚元素"""
        elements = []
        position = start_position
        
        for section in doc.sections:
            # 处理页眉
            if section.header:
                for paragraph in section.header.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        elements.append(DocxElement(
                            element_type='header',
                            content=text,
                            section_number=0,
                            position=position,
                            metadata={'source': 'header'}
                        ))
                        position += 1
            
            # 处理页脚
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        elements.append(DocxElement(
                            element_type='footer',
                            content=text,
                            section_number=0,
                            position=position,
                            metadata={'source': 'footer'}
                        ))
                        position += 1
        
        return elements
    
    def _classify_paragraph_type(self, paragraph: Paragraph) -> str:
        """分类段落类型"""
        style_name = paragraph.style.name.lower()
        
        if 'heading' in style_name:
            return 'header'
        elif 'title' in style_name:
            return 'title'
        elif 'list' in style_name or 'bullet' in style_name:
            return 'list'
        else:
            return 'text'
    
    def _table_to_structured_text(self, table: Table) -> str:
        """将表格转换为结构化文本"""
        try:
            table_data = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_data.append(cell_text)
                
                if any(cell.strip() for cell in row_data):  # 跳过空行
                    # 第一行作为表头
                    if row_idx == 0:
                        table_data.append('| ' + ' | '.join(row_data) + ' |')
                        table_data.append('|' + '---|' * len(row_data))
                    else:
                        table_data.append('| ' + ' | '.join(row_data) + ' |')
            
            return '\n'.join(table_data)
            
        except Exception as e:
            return f"[表格解析错误: {str(e)}]"
    
    def _get_font_size(self, paragraph: Paragraph) -> Optional[float]:
        """获取段落字体大小"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].font.size.pt if paragraph.runs[0].font.size else None
        except:
            pass
        return None
    
    def _is_bold(self, paragraph: Paragraph) -> bool:
        """检查段落是否加粗"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].bold or False
        except:
            pass
        return False
    
    def _is_italic(self, paragraph: Paragraph) -> bool:
        """检查段落是否斜体"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].italic or False
        except:
            pass
        return False
    
    def _has_header_row(self, table: Table) -> bool:
        """检查表格是否有表头行"""
        try:
            if table.rows:
                first_row = table.rows[0]
                # 简单判断：如果第一行的字体加粗，认为是表头
                for cell in first_row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.runs and paragraph.runs[0].bold:
                            return True
        except:
            pass
        return False